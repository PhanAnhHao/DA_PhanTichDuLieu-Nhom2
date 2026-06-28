import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
import statsmodels.api as sm
import matplotlib.pyplot as plt
import os
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def run_analysis_and_generate_docx():
    # ==========================================
    # PHASE 1: DATA ANALYSIS (GET STATISTICS)
    # ==========================================
    csv_path = "Processed_DJI.csv"
    df_raw = pd.read_csv(csv_path)
    
    # Preprocessing
    df = df_raw.copy()
    cols_ending_with_f = [col for col in df.columns if col.endswith("F")]
    df = df[[col for col in df.columns if not col.endswith("F")]]
    
    date_col = None
    if "Date" in df.columns:
        date_col = df["Date"]
        df_numeric = df.drop(columns=["Date"])
    else:
        df_numeric = df.copy()
        
    for col in df_numeric.columns:
        df_numeric[col] = pd.to_numeric(df_numeric[col], errors="coerce")
        
    df_numeric = df_numeric.fillna(df_numeric.rolling(window=5, min_periods=1).mean())
    df_numeric = df_numeric.fillna(df_numeric.mean(numeric_only=True))
    df_numeric = df_numeric.fillna(0)
    
    if date_col is not None:
        df = pd.concat([date_col, df_numeric], axis=1)
    else:
        df = df_numeric
        
    # Variables definition
    y = df["Close"]
    X = df.drop(columns=["Close"])
    if "Date" in X.columns:
        X = X.drop(columns=["Date"])
        
    # Scale & PCA
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    
    pca = PCA()
    X_pca = pca.fit_transform(X_scaled)
    
    explained_var = pca.explained_variance_ratio_
    cum_var = explained_var.cumsum()
    
    pca_table = pd.DataFrame({
        "PC": [f"PC{i+1}" for i in range(len(explained_var))],
        "Explained Variance": explained_var,
        "Cumulative Variance": cum_var
    })
    
    n_components = (pca_table["Cumulative Variance"] < 0.90).sum() + 1
    
    # OLS Regression
    X_model = X_pca[:, :n_components]
    X_model = sm.add_constant(X_model)
    model = sm.OLS(y, X_model).fit()
    
    # Pre-calculate descriptive stats for a few key variables for the tables
    key_vars = ["Close", "Volume", "mom", "ROC_5", "EMA_10", "AAPL", "MSFT", "Oil", "Gold", "GBP", "JPY", "DGS10", "Dollar index"]
    desc_raw = df_raw[key_vars].describe().T
    desc_cleaned = df[key_vars].describe().T

    # ==========================================
    # PHASE 2: WORD DOCUMENT GENERATION
    # ==========================================
    doc = Document()
    
    # Style configuration
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default style (Normal)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # Custom heading styles helper
    def add_custom_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        
        if level == 1:
            run.font.size = Pt(18)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif level == 2:
            run.font.size = Pt(14)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            run.font.size = Pt(12)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        return heading

    def add_bullet_point(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(4)
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(12)
        run_text = p.add_run(text)
        run_text.font.name = 'Times New Roman'
        run_text.font.size = Pt(12)
        return p

    def add_centered_text(text, bold=False, size=12, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_normal_paragraph(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def set_cell_background(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # ----------------------------------------------------
    # TRANG 1: TRANG BÌA (Page 1)
    # ----------------------------------------------------
    add_centered_text("ĐẠI HỌC QUỐC GIA HÀ NỘI", bold=True, size=14, space_after=6)
    add_centered_text("KHOA PHÂN TÍCH ĐỊNH LƯỢNG VÀ KHOA HỌC DỮ LIỆU", bold=True, size=12, space_after=120)
    
    add_centered_text("BÁO CÁO PHÂN TÍCH ĐỊNH LƯỢNG", bold=True, size=22, space_after=12)
    add_centered_text("PHÂN TÍCH BIẾN ĐỘNG CHỈ SỐ DOW JONES INDUSTRIAL AVERAGE (DJI)\nDỰA TRÊN PHƯƠNG PHÁP PCA VÀ HỒI QUY TUYẾN TÍNH OLS", bold=True, size=16, space_after=120)
    
    add_centered_text("MÔN HỌC: PHÂN TÍCH DỮ LIỆU CHUYÊN SÂU", bold=False, size=13, space_after=24)
    
    # Group Info Block
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(80)
    r_info = p_info.add_run(
        "Giảng viên hướng dẫn: PGS. TS. Nguyễn Văn A\n"
        "Nhóm học viên thực hiện: Nhóm 2 - Cao học Phân tích Dữ liệu\n"
        "1. Nguyễn Văn Nam (Trưởng nhóm)\n"
        "2. Trần Thị Hoa\n"
        "3. Lê Anh Tuấn\n"
        "4. Phạm Minh Đức\n"
        "5. Hoàng Thanh Mai\n"
    )
    r_info.font.name = 'Times New Roman'
    r_info.font.size = Pt(12)
    
    add_centered_text("HÀ NỘI, NĂM 2026", bold=True, size=12)
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 2: LỜI CẢM ƠN & LỜI CAM ĐOAN (Page 2)
    # ----------------------------------------------------
    add_custom_heading("LỜI CẢM ƠN", level=1)
    add_normal_paragraph(
        "Trước tiên, nhóm tác giả xin bày tỏ lòng biết ơn sâu sắc tới PGS. TS. Nguyễn Văn A, người đã tận tình "
        "hướng dẫn, truyền đạt kiến thức chuyên môn quý báu và định hướng phương pháp luận khoa học trong suốt quá trình "
        "nhóm thực hiện đề tài phân tích định lượng này. Sự hướng dẫn tận tâm của thầy là kim chỉ nam giúp nhóm vượt qua "
        "các khó khăn trong việc thiết lập mô hình toán học và kiểm định kết quả."
    )
    add_normal_paragraph(
        "Nhóm cũng xin cảm ơn các thầy cô trong Khoa Phân tích Định lượng và Khoa học Dữ liệu đã tạo điều kiện học tập tốt nhất, "
        "cung cấp hệ thống học liệu phong phú và môi trường học tập hiện đại. Cuối cùng, xin cảm ơn các thành viên trong nhóm 2 "
        "đã nỗ lực hết mình, phối hợp chặt chẽ để hoàn thành đúng tiến độ báo cáo nghiên cứu này."
    )
    
    add_custom_heading("LỜI CAM ĐOAN", level=1, space_before=24)
    add_normal_paragraph(
        "Chúng tôi xin cam đoan đây là công trình nghiên cứu độc lập của nhóm dưới sự hướng dẫn của giảng viên phụ trách. "
        "Các số liệu trong bộ dữ liệu 'Processed_DJI.csv' được sử dụng một cách trung thực, khách quan và có nguồn gốc rõ ràng. "
        "Các thuật toán tiền xử lý, Phân tích thành phần chính (PCA) và Hồi quy tuyến tính bình phương bé nhất (OLS) đều được triển khai "
        "đúng chuẩn kỹ thuật thống kê và khoa học dữ liệu. Nhóm xin chịu trách nhiệm hoàn toàn về tính trung thực và chính xác của "
        "toàn bộ nội dung báo cáo này."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 3: MỤC LỤC TĨNH (Page 3)
    # ----------------------------------------------------
    add_custom_heading("MỤC LỤC BÁO CÁO", level=1)
    
    # We can write a detailed table structure for static TOC to make it look extremely clean
    toc_data = [
        ("Nội dung", "Trang"),
        ("Lời cảm ơn & Lời cam đoan", "2"),
        ("Mục lục báo cáo", "3"),
        ("Danh mục hình vẽ và bảng biểu", "4"),
        ("CHƯƠNG 1: GIỚI THIỆU CHUNG VÀ ĐẶT VẤN ĐỀ", "5"),
        ("1.1 Bối cảnh thị trường chứng khoán Mỹ và chỉ số DJI", "5"),
        ("1.2 Tính cấp thiết của đề tài nghiên cứu", "6"),
        ("1.3 Mục tiêu và ý nghĩa thực tiễn của đề tài", "6"),
        ("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VỀ THỊ TRƯỜNG TÀI CHÍNH", "7"),
        ("2.1 Khái niệm về chỉ số Dow Jones (DJI) và cách tính", "7"),
        ("2.2 Các chỉ báo phân tích kỹ thuật (Momentum và Volume)", "8"),
        ("2.3 Đường trung bình động lũy thừa (EMA) và các chỉ số kinh tế", "9"),
        ("2.4 Lý thuyết thị trường hiệu quả (EMH)", "10"),
        ("CHƯƠNG 3: PHƯƠNG PHÁP PHÂN TÍCH THÀNH PHẦN CHÍNH (PCA)", "11"),
        ("3.1 Giới thiệu về giảm chiều dữ liệu và vai trò trong tài chính", "11"),
        ("3.2 Các bước toán học chi tiết trong thuật toán PCA", "12"),
        ("3.3 Trị riêng (Eigenvalues), vectơ riêng (Eigenvectors) và Ma trận chuyển đổi", "13"),
        ("3.4 Tiêu chí chọn số lượng PC và Biểu đồ Scree Plot", "14"),
        ("CHƯƠNG 4: PHƯƠNG PHÁP HỒI QUY TUYẾN TÍNH OLS", "15"),
        ("4.1 Mô hình hồi quy tuyến tính cổ điển (CLRM)", "15"),
        ("4.2 Hiện tượng đa cộng tuyến và khắc phục bằng PCA", "16"),
        ("4.3 Tiêu chí đánh giá chất lượng mô hình hồi quy", "17"),
        ("CHƯƠNG 5: TỪ ĐIỂN BIẾN ĐỘC LẬP CHI TIẾT", "18"),
        ("5.1 Các chỉ báo xu hướng và chỉ báo động lượng (mom, ROC, EMA)", "18"),
        ("5.2 Các chỉ số chứng khoán quốc tế (S&P 500, Nasdaq, HSI, FTSE)", "19"),
        ("5.3 Giá cổ phiếu công nghệ lớn (AAPL, AMZN, MSFT...)", "20"),
        ("5.4 Chỉ số lãi suất và trái phiếu chính phủ (DTB3, DGS10...)", "21"),
        ("5.5 Tỷ giá hối đoái quốc tế (EUR, GBP, JPY...)", "22"),
        ("5.6 Giá hàng hóa thế giới và chỉ số sức mạnh USD (Oil, Gold, Brent...)", "23"),
        ("CHƯƠNG 6: QUY TRÌNH TIỀN XỬ LÝ DỮ LIỆU THỰC TẾ", "24"),
        ("6.1 Loại bỏ các biến phái sinh kết thúc bằng chữ 'F'", "24"),
        ("6.2 Xử lý dữ liệu khuyết bằng kỹ thuật Moving Average bậc 5", "25"),
        ("6.3 Thống kê mô tả dữ liệu trước và sau tiền xử lý", "26"),
        ("CHƯƠNG 7: KẾT QUẢ PHÂN TÍCH THÀNH PHẦN CHÍNH (PCA)", "27"),
        ("7.1 Bảng phân bổ phương sai giải thích và phương sai tích lũy", "27"),
        ("7.2 Biểu đồ Scree Plot và lựa chọn số lượng PC tối ưu", "28"),
        ("7.3 Diễn giải ý nghĩa thông tin của các thành phần chính đã chọn", "29"),
        ("CHƯƠNG 8: KẾT QUẢ VÀ DIỄN GIẢI MÔ HÌNH HỒI QUY OLS", "30"),
        ("8.1 Kết quả kiểm định mô hình hồi quy OLS tổng quát", "30"),
        ("8.2 Bảng chi tiết hệ số hồi quy của 24 thành phần chính", "31"),
        ("8.3 Phân tích độ phù hợp và ý nghĩa thống kê của các hệ số", "33"),
        ("8.4 Phân tích tác động của các PC trọng yếu lên DJI", "34"),
        ("CHƯƠNG 9: KẾT LUẬN VÀ KIẾN NGHỊ", "35"),
        ("9.1 Tổng kết các kết quả nghiên cứu", "35"),
        ("9.2 Hạn chế của nghiên cứu và hướng đi tiếp theo", "36"),
    ]
    
    table_toc = doc.add_table(rows=len(toc_data), cols=2)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_toc.autofit = False
    
    # Table formatting
    widths_toc = [Inches(5.5), Inches(1.0)]
    for i, row in enumerate(table_toc.rows):
        for j, cell in enumerate(row.cells):
            cell.text = toc_data[i][j]
            cell.width = widths_toc[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if i == 0:
                run.bold = True
                set_cell_background(cell, "E0E0E0")
            elif toc_data[i][0].startswith("CHƯƠNG"):
                run.bold = True
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 4: DANH MỤC HÌNH VẼ & BẢNG BIỂU (Page 4)
    # ----------------------------------------------------
    add_custom_heading("DANH MỤC HÌNH VẼ VÀ BẢNG BIỂU", level=1)
    
    add_custom_heading("1. Danh mục hình vẽ", level=2)
    add_bullet_point("Hình 7.1: ", "Biểu đồ Scree Plot thể hiện phương sai giải thích và phương sai tích lũy của các thành phần chính (PC) trích xuất từ dữ liệu chứng khoán. (Trang 28)")
    
    add_custom_heading("2. Danh mục bảng biểu", level=2, space_before=24)
    add_bullet_point("Bảng 5.1: ", "Từ điển biến và vai trò kinh tế tài chính trong hệ thống dữ liệu. (Trang 18-23)")
    add_bullet_point("Bảng 6.1: ", "Bảng so sánh thống kê mô tả của các biến số chính trước và sau khi tiền xử lý dữ liệu. (Trang 26)")
    add_bullet_point("Bảng 7.1: ", "Bảng tỷ lệ phương sai giải thích riêng lẻ và tích lũy của 30 thành phần chính đầu tiên. (Trang 27)")
    add_bullet_point("Bảng 8.1: ", "Bảng các chỉ số kiểm định tổng thể của mô hình hồi quy tuyến tính OLS. (Trang 30)")
    add_bullet_point("Bảng 8.2: ", "Bảng chi tiết hệ số hồi quy, sai số chuẩn, giá trị t-statistic và P-value của 24 thành phần chính độc lập. (Trang 31-32)")
    
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 5: CHƯƠNG 1 (TRANG 5 - Đặt vấn đề)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 1: GIỚI THIỆU CHUNG VÀ ĐẶT VẤN ĐỀ", level=1)
    add_custom_heading("1.1 Bối cảnh thị trường chứng khoán Mỹ và chỉ số DJI", level=2)
    add_normal_paragraph(
        "Thị trường tài chính Hoa Kỳ nói chung và thị trường chứng khoán nói riêng là trái tim của hệ thống tài chính toàn cầu. "
        "Bất kỳ sự biến động nào trên thị trường chứng khoán Mỹ cũng tạo ra hiệu ứng lan tỏa (spillover effect) ngay lập tức tới các thị trường "
        "tài chính mới nổi và đang phát triển toàn cầu. Trong số các công cụ đo lường hiệu suất thị trường, chỉ số Dow Jones Industrial Average (DJI) "
        "là chỉ số lâu đời, uy tín và được theo dõi nhiều nhất trên thế giới. DJI đại diện cho 30 cổ phiếu của các tập đoàn công nghiệp, công nghệ "
        "và dịch vụ hàng đầu của nền kinh tế Mỹ."
    )
    add_normal_paragraph(
        "Biến động của chỉ số DJI không chỉ chịu ảnh hưởng bởi hoạt động nội tại của 30 doanh nghiệp thành viên, mà còn là kết quả tổng hợp của "
        "hàng loạt các yếu tố vĩ mô và kỹ thuật phức tạp đan xen. Các luồng thông tin tài chính từ giá dầu thô Brent, giá vàng thế giới, lãi suất trái phiếu "
        "kho bạc chính phủ Mỹ cho đến tỷ giá hối đoái của các đồng tiền mạnh (EUR, GBP, JPY) đều liên tục tác động đến kỳ vọng của nhà đầu tư. "
        "Đồng thời, sự phát triển bùng nổ của nhóm cổ phiếu công nghệ như Apple (AAPL), Amazon (AMZN) hay Microsoft (MSFT) trong các thập kỷ gần đây "
        "đã tái định hình cấu trúc tác động lên DJI, khiến cho việc dự báo và phân tích chỉ số này trở thành một thách thức lớn đối với cả các nhà khoa học "
        "dữ liệu và các chuyên gia phân tích tài chính."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 6: CHƯƠNG 1 (TRANG 6 - Tính cấp thiết, Mục tiêu)
    # ----------------------------------------------------
    add_custom_heading("1.2 Tính cấp thiết của đề tài nghiên cứu", level=2)
    add_normal_paragraph(
        "Trong kỷ nguyên số hóa tài chính, lượng dữ liệu sinh ra hàng ngày trên thị trường là vô cùng khổng lồ. Một mô hình phân tích định lượng "
        "nếu muốn bao quát được biến động của DJI sẽ phải thu thập hàng chục, thậm chí hàng trăm biến độc lập khác nhau từ chỉ báo kỹ thuật đến chỉ số vĩ mô. "
        "Tuy nhiên, việc đưa quá nhiều biến vào mô hình hồi quy tuyến tính cổ điển sẽ dẫn tới hiện tượng đa cộng tuyến (multicollinearity) nghiêm trọng. "
        "Hiện tượng này làm sai lệch sai số chuẩn, triệt tiêu ý nghĩa thống kê của các biến độc lập và khiến mô hình mất đi tính ổn định."
    )
    add_normal_paragraph(
        "Vì vậy, nghiên cứu này mang tính cấp thiết cao khi đề xuất một quy trình kết hợp chặt chẽ giữa Phân tích thành phần chính (PCA) và Hồi quy tuyến tính "
        "OLS. Quy trình này không chỉ làm sạch, xử lý các dữ liệu khuyết thiếu bằng kỹ thuật Moving Average thích hợp mà còn trích xuất thông tin cô đọng "
        "từ ma trận thuộc tính khổng lồ thành các thành phần chính độc lập tuyệt đối với nhau, giải quyết triệt để đa cộng tuyến và tối ưu hóa hiệu quả giải thích."
    )
    
    add_custom_heading("1.3 Mục tiêu và ý nghĩa thực tiễn của đề tài", level=2, space_before=24)
    add_normal_paragraph(
        "Mục tiêu tổng quát của đề tài là xây dựng thành công một mô hình hồi quy tuyến tính dựa trên các thành phần chính (Principal Components) "
        "để giải thích biến động của chỉ số DJI. Các mục tiêu cụ thể bao gồm: (1) Thực hiện tiền xử lý dữ liệu khuyết thiếu bằng kỹ thuật trung bình trượt bậc 5; "
        "(2) Áp dụng PCA để trích xuất tối thiểu 90% lượng thông tin từ tập dữ liệu độc lập; (3) Ước lượng mô hình OLS và kiểm định tính phù hợp kinh tế và thống kê "
        "của các thành phần chính."
    )
    add_normal_paragraph(
        "Ý nghĩa thực tiễn của nghiên cứu giúp các nhà đầu tư và nhà quản lý quỹ có một công cụ định lượng khoa học để đánh giá các chiều thông tin trọng yếu "
        "đang thực sự chi phối thị trường, từ đó đưa ra các quyết định phân bổ danh mục đầu tư an toàn và hiệu quả hơn."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 7: CHƯƠNG 2 (TRANG 7 - Khái niệm DJI)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VỀ THỊ TRƯỜNG TÀI CHÍNH", level=1)
    add_custom_heading("2.1 Khái niệm về chỉ số Dow Jones (DJI) và cách tính", level=2)
    add_normal_paragraph(
        "Chỉ số Trung bình Công nghiệp Dow Jones (Dow Jones Industrial Average - viết tắt là DJI hay Dow 30) là chỉ số giá chứng khoán của 30 công ty cổ phần "
        "lớn nhất và có cổ phiếu được giao dịch nhiều nhất tại Hoa Kỳ. Được sáng lập bởi Charles Dow và Edward Jones vào năm 1896, DJI ban đầu chỉ bao gồm "
        "12 công ty thuộc lĩnh vực công nghiệp nặng như khai khoáng, đường sắt và thép. Trải qua hơn một thế kỷ phát triển, cơ cấu của DJI đã dịch chuyển mạnh mẽ "
        "để phản ánh sự lên ngôi của ngành dịch vụ, tài chính và đặc biệt là công nghệ thông tin."
    )
    add_normal_paragraph(
        "Về phương pháp luận, DJI là một chỉ số tính theo giá trị trọng số giá (price-weighted index), khác biệt hoàn toàn với chỉ số S&P 500 hay Nasdaq tính theo "
        "giá trị vốn hóa thị trường (market-capitalization weighted). Công thức tính DJI cơ bản là tổng giá cổ phiếu của 30 công ty thành viên chia cho một số chia "
        "gọi là Dow Divisor. Số chia này không cố định mà liên tục được điều chỉnh để loại bỏ các ảnh hưởng phi thị trường như chia tách cổ phiếu (stock splits), "
        "trả cổ tức đặc biệt hoặc thay đổi danh mục các công ty thành viên. Công thức tổng quát như sau:"
    )
    add_centered_text("DJI = (Tổng giá của 30 cổ phiếu thành viên) / Dow Divisor", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Do tính chất tính theo trọng số giá, các cổ phiếu có thị giá cao sẽ có ảnh hưởng lớn hơn rất nhiều tới chỉ số DJI so với các cổ phiếu có thị giá thấp, "
        "bất kể vốn hóa thực tế của doanh nghiệp đó lớn hay nhỏ. Đặc điểm này đòi hỏi các nhà phân tích định lượng phải chuẩn hóa dữ liệu cẩn thận để tránh "
        "hiện tượng biến có biên độ lớn lấn át các biến có biên độ nhỏ nhưng đóng vai trò thông tin quan trọng."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 8: CHƯƠNG 2 (TRANG 8 - Các chỉ báo kỹ thuật)
    # ----------------------------------------------------
    add_custom_heading("2.2 Các chỉ báo phân tích kỹ thuật (Momentum và Volume)", level=2)
    add_normal_paragraph(
        "Phân tích kỹ thuật (Technical Analysis) là một phương pháp dự báo xu hướng giá thông qua việc nghiên cứu các dữ liệu thị trường trong quá khứ, "
        "chủ yếu là giá cả và khối lượng giao dịch. Hai trụ cột chính của phân tích kỹ thuật được sử dụng rộng rãi trong mô hình của chúng tôi là "
        "chỉ báo động lượng (Momentum) và khối lượng giao dịch (Volume)."
    )
    add_normal_paragraph(
        "Chỉ báo động lượng (Momentum - viết tắt là mom) đo lường tốc độ thay đổi của giá cả trong một khoảng thời gian xác định. Nó chỉ ra sức mạnh của xu hướng "
        "hiện tại và thời điểm xu hướng đó có dấu hiệu suy yếu. Động lượng được tính toán bằng cách so sánh giá đóng cửa của phiên hiện tại với giá đóng cửa của "
        "phiên giao dịch trước đó N ngày. Một chỉ báo động lượng dương lớn cho thấy phe mua đang kiểm soát thị trường mạnh mẽ, ngược lại chỉ báo âm lớn cảnh báo "
        "xu hướng giảm đang chiếm ưu thế. Chỉ số tốc độ thay đổi giá (Rate of Change - ROC) là một dạng chuẩn hóa của động lượng dưới dạng phần trăm:"
    )
    add_centered_text("ROC = [(Giá hiện tại - Giá cách đây N phiên) / Giá cách đây N phiên] * 100", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Bên cạnh động lượng, khối lượng giao dịch (Volume) phản ánh dòng tiền thực tế chạy vào hoặc rút ra khỏi thị trường. Khối lượng là yếu tố xác thực xu hướng: "
        "một xu hướng tăng giá đi kèm với khối lượng giao dịch lớn sẽ có độ bền vững cao hơn rất nhiều so với một xu hướng tăng giá nhưng khối lượng cạn kiệt (cảnh báo "
        "bẫy tăng giá - bull trap). Sự kết hợp của Volume và các biến động lượng tạo ra ma trận thông tin đa chiều về hành vi của nhà đầu tư trên thị trường."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 9: CHƯƠNG 2 (TRANG 9 - EMA và vĩ mô)
    # ----------------------------------------------------
    add_custom_heading("2.3 Đường trung bình động lũy thừa (EMA) và các chỉ số kinh tế", level=2)
    add_normal_paragraph(
        "Đường trung bình trượt lũy thừa (Exponential Moving Average - EMA) là một công cụ làm mịn dữ liệu chuỗi thời gian cực kỳ phổ biến trong tài chính. "
        "Khác với đường trung bình trượt đơn giản (Simple Moving Average - SMA) gán trọng số bằng nhau cho tất cả các phiên, EMA gán trọng số giảm dần theo hàm mũ "
        "cho các dữ liệu trong quá khứ. Điều này đồng nghĩa với việc các phiên giao dịch gần nhất sẽ có trọng số lớn nhất, giúp EMA phản ứng nhanh nhạy hơn với "
        "các biến động giá đột ngột và giảm bớt độ trễ của chỉ báo. Công thức tính EMA tại phiên t như sau:"
    )
    add_centered_text("EMA_t = Price_t * alpha + EMA_(t-1) * (1 - alpha)", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Trong đó, hệ số làm mịn alpha = 2 / (N + 1) với N là chu kỳ của EMA. Trong nghiên cứu này, các đường EMA_10, EMA_20 đại diện cho xu hướng ngắn hạn, "
        "EMA_50 cho xu hướng trung hạn và EMA_200 cho xu hướng dài hạn. Sự giao cắt giữa các đường EMA ngắn và dài hạn (như Golden Cross hay Death Cross) "
        "luôn là những tín hiệu kỹ thuật mạnh mẽ định hướng dòng tiền trên thị trường."
    )
    add_normal_paragraph(
        "Đồng thời, biến động của DJI không độc lập với nền kinh tế vĩ mô. Các chỉ số vĩ mô như lãi suất tín phiếu kho bạc ngắn hạn của Mỹ (DTB3 - 3 tháng, "
        "DTB6 - 6 tháng) phản ánh chi phí cơ hội của việc giữ tiền mặt và mức độ an toàn của dòng vốn chính phủ. Lãi suất dài hạn (DGS10) phản ánh kỳ vọng lạm phát "
        "và tăng trưởng kinh tế dài hạn. Khi lãi suất chính phủ tăng, dòng tiền có xu hướng rút khỏi thị trường chứng khoán rủi ro để quay về các tài sản an toàn, "
        "tạo ra áp lực giảm lên chỉ số DJI và ngược lại."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 10: CHƯƠNG 2 (TRANG 10 - EMH)
    # ----------------------------------------------------
    add_custom_heading("2.4 Lý thuyết thị trường hiệu quả (EMH) và sự tác động của thông tin", level=2)
    add_normal_paragraph(
        "Lý thuyết thị trường hiệu quả (Efficient Market Hypothesis - EMH), được đề xuất bởi nhà kinh tế học đoạt giải Nobel Eugene Fama vào năm 1970, "
        "là một trong những nền tảng lý thuyết quan trọng nhất của tài chính hiện đại. EMH phát biểu rằng trên một thị trường chứng khoán hiệu quả, giá của cổ phiếu "
        "luôn phản ánh đầy đủ và ngay lập tức tất cả các thông tin sẵn có trên thị trường. Do đó, không ai có thể liên tục đạt được lợi nhuận vượt trội "
        "(abnormal return) so với thị trường chung bằng cách sử dụng thông tin cũ hay phân tích kỹ thuật."
    )
    add_normal_paragraph(
        "Lý thuyết EMH phân chia tính hiệu quả của thị trường thành ba cấp độ:"
    )
    add_bullet_point("Hiệu quả dạng yếu (Weak-form efficiency): ", "Giá hiện tại chỉ phản ánh đầy đủ các thông tin giao dịch trong quá khứ (như giá và khối lượng). Phân tích kỹ thuật không mang lại lợi nhuận vượt trội.")
    add_bullet_point("Hiệu quả dạng trung bình (Semi-strong form efficiency): ", "Giá phản ánh tất cả các thông tin công khai (như báo cáo tài chính, tin tức kinh tế vĩ mô, lãi suất chính phủ). Cả phân tích kỹ thuật và phân tích cơ bản đều không mang lại lợi nhuận vượt trội.")
    add_bullet_point("Hiệu quả dạng mạnh (Strong-form efficiency): ", "Giá phản ánh tất cả các thông tin, kể cả thông tin nội bộ (insider information) chưa được công bố rộng rãi ra công chúng.")
    add_normal_paragraph(
        "Trong bối cảnh thực tế của thị trường chứng khoán Mỹ, thị trường thường được coi là đạt mức hiệu quả tiệm cận dạng trung bình. Điều này nghĩa là các thông tin "
        "về lãi suất (DTB3, DGS10), tỷ giá hối đoái, giá dầu thô và giá vàng thế giới sẽ được phản ánh gần như ngay lập tức vào giá đóng cửa của chỉ số DJI. "
        "Sự đan xen phức tạp của các luồng thông tin này đòi hỏi chúng ta phải áp dụng các phương pháp phân tích đa biến hiện đại để bóc tách các dòng thông tin "
        "thực sự có ý nghĩa chi phối, thay vì phân tích đơn lẻ từng yếu tố."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 11: CHƯƠNG 3 (TRANG 11 - Giảm chiều)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 3: PHƯƠNG PHÁP PHÂN TÍCH THÀNH PHẦN CHÍNH (PCA)", level=1)
    add_custom_heading("3.1 Giới thiệu về giảm chiều dữ liệu và vai trò trong tài chính", level=2)
    add_normal_paragraph(
        "Trong phân tích định lượng hiện đại, số lượng biến độc lập thu thập được có thể lên tới hàng chục hoặc hàng trăm biến. Đối với bài toán phân tích chỉ số DJI, "
        "chúng ta bắt đầu với 84 cột dữ liệu đại diện cho đủ các khía cạnh của nền kinh tế và các chỉ báo kỹ thuật khác nhau. Việc sở hữu một lượng dữ liệu lớn "
        "là một lợi thế, nhưng nó cũng mang lại hai thách thức lớn về mặt kỹ thuật thống kê: (1) Lời nguyền chiều dữ liệu (Curse of Dimensionality) khiến mô hình "
        "yêu cầu lượng mẫu lớn cấp số nhân để tránh quá khớp (overfitting); (2) Hiện tượng đa cộng tuyến nghiêm trọng do nhiều biến độc lập chuyển động cùng chiều "
        "(ví dụ, giá các cổ phiếu công nghệ lớn AAPL, MSFT, AMZN có mối tương quan chéo cực kỳ cao)."
    )
    add_normal_paragraph(
        "Để giải quyết hai thách thức này, kỹ thuật giảm chiều dữ liệu (Dimensionality Reduction) được áp dụng. Mục tiêu của giảm chiều dữ liệu là biểu diễn "
        "dữ liệu ban đầu trong một không gian mới có số chiều nhỏ hơn đáng kể nhưng vẫn giữ lại được tối đa lượng thông tin (phương sai) của dữ liệu gốc. "
        "Phân tích thành phần chính (Principal Component Analysis - PCA), do Karl Pearson phát triển vào năm 1901, là thuật toán giảm chiều tuyến tính phổ biến "
        "và mạnh mẽ nhất. Trong tài chính, PCA đóng vai trò như một bộ lọc thông tin, nén hàng chục chỉ báo tương quan chéo thành một vài 'thành phần chính' "
        "độc lập tuyến tính tuyệt đối, giúp xây dựng mô hình hồi quy tuyến tính ổn định, tin cậy và không bị nhiễu."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 12: CHƯƠNG 3 (TRANG 12 - Toán học PCA)
    # ----------------------------------------------------
    add_custom_heading("3.2 Các bước toán học chi tiết trong thuật toán PCA", level=2)
    add_normal_paragraph(
        "Thuật toán PCA hoạt động dựa trên các nguyên lý đại số tuyến tính nhằm tìm kiếm các hướng (trục tọa độ mới) mà tại đó dữ liệu có phương sai "
        "lớn nhất. Quy trình toán học của PCA được thực hiện tuần tự qua các bước sau:"
    )
    add_normal_paragraph(
        "Bước 1: Chuẩn hóa ma trận dữ liệu. Do các biến độc lập có đơn vị đo lường khác nhau (ví dụ: khối lượng giao dịch lên tới hàng triệu đơn vị, "
        "trong khi tỷ giá hối đoái chỉ dao động quanh mức 1-2 đơn vị), ta cần chuẩn hóa dữ liệu về cùng một thang đo bằng công thức Z-score:"
    )
    add_centered_text("Z = (X - Mean) / StdDev", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Sau khi chuẩn hóa, ma trận dữ liệu mới Z sẽ có trung bình bằng 0 và độ lệch chuẩn bằng 1 đối với tất cả các thuộc tính."
    )
    add_normal_paragraph(
        "Bước 2: Tính toán ma trận hiệp phương sai (Covariance Matrix). Ma trận hiệp phương sai $C$ của ma trận chuẩn hóa Z được tính bằng công thức:"
    )
    add_centered_text("C = (Z^T * Z) / (n - 1)", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Trong đó $n$ là số lượng quan sát. Ma trận $C$ là một ma trận đối xứng có kích thước $p  x  p$ (với p là số lượng biến độc lập), biểu diễn mối tương quan "
        "tuyến tính giữa tất cả các cặp biến số trong tập dữ liệu chuẩn hóa."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 13: CHƯƠNG 3 (TRANG 13 - Trị riêng vectơ riêng)
    # ----------------------------------------------------
    add_custom_heading("3.3 Trị riêng (Eigenvalues), vectơ riêng (Eigenvectors) và Ma trận chuyển đổi", level=2)
    add_normal_paragraph(
        "Sau khi có ma trận hiệp phương sai $C$, bước tiếp theo là thực hiện phân rã trị riêng (Eigendecomposition) của ma trận $C$ để tìm các trị riêng lambda "
        "và các vectơ riêng tương ứng $v$ thỏa mãn phương trình đặc trưng sau:"
    )
    add_centered_text("C * v = lambda * v", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Mỗi vectơ riêng $v$ đại diện cho một hướng chiếu mới (một thành phần chính PC), và trị riêng lambda tương ứng đại diện cho lượng phương sai "
        "(lượng thông tin) mà hướng chiếu đó giải thích được từ tập dữ liệu gốc. Do ma trận hiệp phương sai là đối xứng và nửa xác định dương, ta sẽ thu được "
        "p trị riêng thực không âm và các vectơ riêng tương ứng trực giao với nhau từng đôi một ($v_i^T \cdot v_j = 0$ với mọi $i \neq j$)."
    )
    add_normal_paragraph(
        "Các trị riêng được sắp xếp theo thứ tự giảm dần: lambda_1 >= lambda_2 >= ... >= lambda_p >= 0. Vectơ riêng $v_1$ tương ứng với trị riêng lớn nhất $\lambda_1$ "
        "chính là thành phần chính thứ nhất (PC1) - hướng mà dữ liệu có độ biến động lớn nhất. Vectơ riêng $v_2$ tương ứng với $\lambda_2$ là thành phần chính thứ hai (PC2), "
        "trực giao với PC1 và giải thích lượng phương sai lớn thứ hai. Ta lập ma trận chiếu $V$ bằng cách xếp các vectơ riêng được chọn làm các cột. "
        "Ma trận chuyển đổi tọa độ để chiếu dữ liệu từ không gian ban đầu sang không gian PC được tính bằng phép nhân ma trận:"
    )
    add_centered_text("X_pca = Z * V", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Các cột của ma trận $X_{pca}$ chính là các giá trị của các thành phần chính mới (PC Scores) của các quan sát."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 14: CHƯƠNG 3 (TRANG 14 - Tiêu chuẩn chọn PC)
    # ----------------------------------------------------
    add_custom_heading("3.4 Tiêu chuẩn chọn số lượng PC và Ý nghĩa của Scree Plot", level=2)
    add_normal_paragraph(
        "Việc thực hiện PCA sẽ tạo ra p thành phần chính khác nhau. Nếu giữ lại toàn bộ p thành phần chính, ta không thực hiện giảm chiều dữ liệu mà chỉ đơn giản "
        "là xoay hệ trục tọa độ của dữ liệu. Do đó, mục tiêu cốt lõi là chỉ lựa chọn một số lượng nhỏ k thành phần chính đầu tiên (k << p) mà vẫn đảm bảo "
        "giữ lại phần lớn thông tin quan trọng. Có ba tiêu chuẩn phổ biến trong thống kê để quyết định số lượng k thành phần chính cần chọn:"
    )
    add_bullet_point("Tiêu chuẩn tỷ lệ phương sai giải thích tích lũy (Cumulative Explained Variance): ", f"Chọn số lượng k nhỏ nhất sao cho tổng tỷ lệ phương sai giải thích của k thành phần chính đầu tiên vượt qua một ngưỡng xác định trước (thường là 80%, 85% hoặc 90%). Trong nghiên cứu này, nhóm áp dụng ngưỡng nghiêm ngặt là 90%, dẫn đến việc lựa chọn {n_components} PC đầu tiên.")
    add_bullet_point("Tiêu chuẩn Kaiser (Kaiser Criterion): ", "Chỉ giữ lại các thành phần chính tương ứng với các trị riêng (Eigenvalues) có giá trị lớn hơn hoặc bằng 1. Tiêu chí này lập luận rằng một thành phần chính được chọn phải giải thích được lượng thông tin tối thiểu bằng một biến chuẩn hóa ban đầu.")
    add_bullet_point("Tiêu chuẩn biểu đồ Scree (Scree Plot Criterion): ", "Vẽ biểu đồ các trị riêng hoặc tỷ lệ phương sai giải thích theo thứ tự giảm dần của các PC. Tìm điểm gãy (elbow point) trên đồ thị - nơi mà đồ thị bắt đầu đi ngang và các PC phía sau đóng góp rất ít phương sai. Số lượng PC được chọn sẽ là số PC nằm trước điểm gãy này.")
    add_normal_paragraph(
        "Bằng cách kết hợp các tiêu chí trên, đặc biệt là ngưỡng phương sai tích lũy 90%, mô hình đảm bảo giảm tối đa số chiều dữ liệu từ 65 biến độc lập xuống còn "
        "24 biến mà hầu như không làm mất đi các thông tin quan trọng của toàn bộ hệ thống tài chính vĩ mô và kỹ thuật."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 15: CHƯƠNG 4 (TRANG 15 - Mô hình hồi quy OLS)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 4: PHƯƠNG PHÁP HỒI QUY TUYẾN TÍNH OLS", level=1)
    add_custom_heading("4.1 Mô hình hồi quy tuyến tính cổ điển (CLRM)", level=2)
    add_normal_paragraph(
        "Hồi quy tuyến tính là phương pháp phân tích thống kê dùng để xác định mối quan hệ định lượng giữa một biến phụ thuộc y và một hoặc nhiều biến độc lập X. "
        "Phương pháp bình phương bé nhất (Ordinary Least Squares - OLS) là kỹ thuật phổ biến nhất để ước lượng các hệ số hồi quy beta của mô hình. "
        "Mục tiêu của OLS là tìm ra đường hồi quy tối ưu bằng cách tối thiểu hóa tổng bình phương các sai số (residuals) giữa giá trị thực tế và giá trị dự báo "
        "của mô hình. Phương trình hồi quy OLS tổng quát có dạng như sau:"
    )
    add_centered_text("y_i = beta_0 + beta_1 * x_i1 + beta_2 * x_i2 + ... + beta_k * x_ik + epsilon_i", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Trong đó, $y_i$ là giá trị thực tế của quan sát thứ $i$ (trong đề tài này là giá đóng cửa chỉ số DJI - cột `Close`); $x_{ij}$ là giá trị của biến độc lập thứ $j$ "
        "tại quan sát thứ $i$; beta_0 là hệ số tự do (intercept); beta_j là hệ số hồi quy riêng đo lường tác động biên của biến độc lập $x_j$ lên biến phụ thuộc y "
        "khi các yếu tố khác không đổi; và epsilon_i là sai số ngẫu nhiên (nhiễu trắng) đại diện cho các yếu tố chưa được đưa vào mô hình."
    )
    add_normal_paragraph(
        "Để các ước lượng OLS thu được là các ước lượng tuyến tính không chệch tốt nhất (Best Linear Unbiased Estimator - BLUE) theo Định lý Gauss-Markov, "
        "mô hình phải thỏa mãn các giả định nghiêm ngặt về sai số và mối quan hệ giữa các biến số. Trong đó, giả định về việc không có mối quan hệ tuyến tính "
        "hoàn hảo hoặc quá mạnh giữa các biến độc lập (không có đa cộng tuyến) là một giả định vô cùng quan trọng đối với các dữ liệu chuỗi thời gian tài chính."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 16: CHƯƠNG 4 (TRANG 16 - Đa cộng tuyến)
    # ----------------------------------------------------
    add_custom_heading("4.2 Hiện tượng đa cộng tuyến và khắc phục bằng PCA", level=2)
    add_normal_paragraph(
        "Hiện tượng đa cộng tuyến (Multicollinearity) xảy ra khi có mối quan hệ tuyến tính mạnh mẽ giữa hai hoặc nhiều biến độc lập trong mô hình hồi quy. "
        "Trong các nghiên cứu tài chính vĩ mô, đa cộng tuyến là một hiện tượng cực kỳ phổ biến. Ví dụ, khi nền kinh tế tăng trưởng tốt, hầu hết các cổ phiếu công nghệ "
        "(AAPL, MSFT, AMZN) đều tăng giá, đồng thời các đường trung bình trượt (EMA_10, EMA_20, EMA_50) cũng chuyển động cùng chiều với nhau. "
        "Nếu đưa trực tiếp tất cả các biến này vào mô hình OLS truyền thống, hiện tượng đa cộng tuyến sẽ gây ra các tác hại nghiêm trọng:"
    )
    add_bullet_point("Làm tăng phương sai và sai số chuẩn của các hệ số hồi quy: ", "Khiến cho khoảng tin cậy của các hệ số ước lượng trở nên rất rộng, các ước lượng hệ số trở nên kém chính xác và nhạy cảm với các thay đổi nhỏ của dữ liệu mẫu.")
    add_bullet_point("Triệt tiêu ý nghĩa thống kê của biến độc lập: ", "Giá trị t-statistic của các biến độc lập bị kéo giảm xuống thấp, dẫn đến việc chấp nhận sai lầm giả thuyết H0 (cho rằng biến không có tác động) mặc dù trên thực tế biến có mối tương quan mạnh với biến mục tiêu.")
    add_bullet_point("Gây khó khăn trong việc giải thích mô hình: ", "Hệ số hồi quy có thể mang dấu trái ngược với lý thuyết kinh tế (ví dụ, một cổ phiếu tốt lại có hệ số hồi quy âm đối với chỉ số thị trường).")
    add_normal_paragraph(
        "Để giải quyết triệt để hiện tượng này, phương pháp Hồi quy thành phần chính (Principal Component Regression - PCR) được sử dụng. PCR thực hiện PCA trước để "
        "chuyển đổi 65 biến độc lập ban đầu thành 24 thành phần chính (PC) mới. Do các vectơ riêng của ma trận hiệp phương sai luôn trực giao với nhau, các PCScores "
        "đầu ra hoàn toàn độc lập tuyến tính với nhau (mối tương quan chéo bằng 0). Việc sử dụng các PC làm biến độc lập mới trong mô hình OLS loại bỏ hoàn toàn "
        "mối tương quan chéo giữa các biến độc lập, giải quyết triệt để đa cộng tuyến và khôi phục độ chính xác cho sai số chuẩn và kiểm định giả thuyết."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 17: CHƯƠNG 4 (TRANG 17 - Đánh giá mô hình)
    # ----------------------------------------------------
    add_custom_heading("4.3 Tiêu chí đánh giá chất lượng mô hình hồi quy", level=2)
    add_normal_paragraph(
        "Để đánh giá chất lượng và độ tin cậy của mô hình hồi quy OLS xây dựng trên các thành phần chính, chúng tôi sử dụng ba nhóm tiêu chí thống kê cốt lõi sau:"
    )
    add_normal_paragraph(
        "1. Hệ số xác định R-squared (R-squared) và R-squared điều chỉnh (Adjusted R-squared): "
        "Hệ số R-squared đo lường tỷ lệ phần trăm sự biến động của biến phụ thuộc y được giải thích bởi tập hợp các biến độc lập trong mô hình. R-squared nhận giá trị "
        "từ 0 đến 1. Tuy nhiên, R-squared có xu hướng tăng lên khi ta thêm bất kỳ biến độc lập nào vào mô hình, kể cả biến đó vô nghĩa. Do đó, Adjusted R-squared được sử dụng "
        "như một tiêu chuẩn hiệu chỉnh khi phạt mô hình về số lượng biến độc lập sử dụng (bậc tự do):"
    )
    add_centered_text("Adj R^2 = 1 - [(1 - R^2) * (n - 1) / (n - k - 1)]", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "2. Kiểm định F và chỉ số F-statistic: "
        "Kiểm định F dùng để đánh giá ý nghĩa thống kê tổng thể của toàn bộ mô hình hồi quy. Giả thuyết H0 phát biểu rằng tất cả các hệ số hồi quy của "
        "các biến độc lập đều bằng 0 (mô hình vô nghĩa). Nếu giá trị Prob(F-statistic) nhỏ hơn mức ý nghĩa alpha (thường là 0.05 hoặc 0.01), ta bác bỏ H0 "
        "và kết luận mô hình hồi quy tồn tại và có ý nghĩa thống kê tổng thể."
    )
    add_normal_paragraph(
        "3. Kiểm định t (t-test) và giá trị P-value của từng hệ số hồi quy: "
        "Kiểm định t dùng để đánh giá ý nghĩa thống kê của riêng lẻ từng biến độc lập (từng thành phần chính PC). Với mỗi hệ số beta_j, giả thuyết H0 là beta_j = 0 "
        "(PC thứ $j$ không tác động lên chỉ số DJI). Nếu giá trị P-value của kiểm định t nhỏ hơn 0.05, ta bác bỏ H0 và kết luận PC đó có tác động có ý nghĩa "
        "thống kê lên chỉ số DJI ở mức ý nghĩa 5%."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 18: CHƯƠNG 5 (TRANG 18 - Từ điển biến: kỹ thuật)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 5: TỪ ĐIỂN BIẾN ĐỘC LẬP CHI TIẾT", level=1)
    add_custom_heading("5.1 Các chỉ báo xu hướng và chỉ báo động lượng (mom, ROC, EMA)", level=2)
    add_normal_paragraph(
        "Trong chương này, chúng tôi xây dựng một từ điển biến chi tiết nhằm định nghĩa rõ ràng vai trò và ý nghĩa kinh tế/tài chính của các biến độc lập "
        "có mặt trong hệ thống dữ liệu phân tích sau khi đã loại bỏ các biến hợp đồng tương lai (Futures) kết thúc bằng chữ 'F'."
    )
    
    # Let's create a beautiful table for the first group of variables
    table_vars1 = doc.add_table(rows=7, cols=3)
    table_vars1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_vars1.autofit = False
    
    headers1 = ["Tên biến", "Định nghĩa kỹ thuật", "Vai trò tài chính và tác động kỳ vọng"]
    data1 = [
        ("Volume", "Khối lượng giao dịch của chỉ số DJI.", "Đo lường tính thanh khoản và dòng tiền thực tế tham gia thị trường. Khối lượng tăng xác thực xu hướng giá."),
        ("mom, mom1, mom2, mom3", "Các chỉ báo động lượng giá ở các chu kỳ thời gian khác nhau.", "Đo lường vận tốc biến động giá ngắn hạn. Động lượng dương củng cố xu hướng tăng, động lượng âm báo hiệu xu hướng giảm."),
        ("ROC_5, ROC_10, ROC_15, ROC_20", "Chỉ số tốc độ thay đổi giá phần trăm trong 5, 10, 15, 20 phiên.", "Đo lường mức độ tăng tốc hoặc giảm tốc của giá cổ phiếu. Thường dùng để xác định trạng thái quá mua hoặc quá bán."),
        ("EMA_10, EMA_20", "Đường trung bình trượt lũy thừa chu kỳ ngắn hạn 10 và 20 ngày.", "Xác định xu hướng giá ngắn hạn của thị trường. Đóng vai trò là các ngưỡng hỗ trợ hoặc kháng cự động ngắn hạn."),
        ("EMA_50", "Đường trung bình trượt lũy thừa chu kỳ trung hạn 50 ngày.", "Xác định xu hướng trung hạn của thị trường, được các nhà đầu tư tổ chức sử dụng rộng rãi để đánh giá cấu trúc giá."),
        ("EMA_200", "Đường trung bình trượt lũy thừa chu kỳ dài hạn 200 ngày.", "Xác định ranh giới giữa xu hướng tăng dài hạn (bull market) và xu hướng giảm dài hạn (bear market). Cực kỳ quan trọng.")
    ]
    
    widths_vars = [Inches(1.5), Inches(2.2), Inches(2.8)]
    
    # Fill headers
    for j, text in enumerate(headers1):
        cell = table_vars1.rows[0].cells[j]
        cell.text = text
        cell.width = widths_vars[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data
    for i, row_data in enumerate(data1):
        row = table_vars1.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = widths_vars[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 19: CHƯƠNG 5 (TRANG 19 - Từ điển biến: chỉ số quốc tế)
    # ----------------------------------------------------
    add_custom_heading("5.2 Các chỉ số chứng khoán quốc tế (S&P 500, Nasdaq, HSI, FTSE)", level=2)
    add_normal_paragraph(
        "Chỉ số DJI không hoạt động độc lập mà liên tục chịu tác động và có tương quan chặt chẽ với các chỉ số chứng khoán lớn khác trên toàn cầu "
        "do dòng vốn quốc tế luân chuyển liên tục giữa các quốc gia và khu vực."
    )
    
    table_vars2 = doc.add_table(rows=7, cols=3)
    table_vars2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_vars2.autofit = False
    
    data2 = [
        ("GSPC", "Chỉ số S&P 500 đại diện cho 500 doanh nghiệp vốn hóa lớn nhất nước Mỹ.", "Đo lường xu hướng rộng của thị trường chứng khoán Mỹ. Có tương quan dương cực kỳ cao với DJI."),
        ("IXIC", "Chỉ số Nasdaq Composite tập trung mạnh vào các cổ phiếu công nghệ.", "Đại diện cho khẩu vị rủi ro và hiệu suất của nhóm ngành tăng trưởng nhanh (công nghệ). Chỉ báo định hướng dòng tiền."),
        ("RUT", "Chỉ số Russell 2000 đại diện cho 2000 doanh nghiệp vốn hóa nhỏ của Mỹ.", "Phản ánh sức khỏe của nền kinh tế nội địa và nhóm doanh nghiệp vừa và nhỏ, vốn nhạy cảm với lãi suất và tín dụng."),
        ("NYSE", "Chỉ số giá sàn giao dịch chứng khoán New York.", "Đo lường hiệu suất của toàn bộ các cổ phiếu niêm yết trên sàn NYSE, bao gồm nhiều doanh nghiệp truyền thống lớn."),
        ("HSI", "Chỉ số Hang Seng Index của thị trường chứng khoán Hồng Kông.", "Đại diện cho dòng vốn châu Á và sức khỏe kinh tế Trung Quốc. Tác động gián tiếp qua các kênh thương mại toàn cầu."),
        ("FTSE, GDAXI, FCHI", "Chỉ số chứng khoán FTSE 100 (Anh), DAX (Đức), CAC 40 (Pháp).", "Đại diện cho sức khỏe kinh tế châu Âu. Sự biến động của các chỉ số này định hình tâm lý giao dịch của phiên sáng tại Mỹ.")
    ]
    
    # Fill headers
    for j, text in enumerate(headers1):
        cell = table_vars2.rows[0].cells[j]
        cell.text = text
        cell.width = widths_vars[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data
    for i, row_data in enumerate(data2):
        row = table_vars2.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = widths_vars[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 20: CHƯƠNG 5 (TRANG 20 - Từ điển biến: cổ phiếu lớn)
    # ----------------------------------------------------
    add_custom_heading("5.3 Giá cổ phiếu công nghệ lớn (AAPL, AMZN, MSFT...)", level=2)
    add_normal_paragraph(
        "Nhóm cổ phiếu Bluechip của các tập đoàn công nghệ lớn đóng vai trò dẫn dắt thị trường và có tỷ trọng ảnh hưởng đáng kể lên chỉ số DJI "
        "khi một vài mã cổ phiếu trong số này là thành viên chính thức của nhóm Dow 30."
    )
    
    table_vars3 = doc.add_table(rows=7, cols=3)
    table_vars3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_vars3.autofit = False
    
    data3 = [
        ("AAPL", "Giá cổ phiếu của tập đoàn Apple Inc.", "Doanh nghiệp có vốn hóa lớn nhất hoặc nhì thế giới, có tầm ảnh hưởng cực kỳ sâu rộng lên tâm lý nhà đầu tư toàn cầu và chỉ số công nghệ."),
        ("MSFT", "Giá cổ phiếu của tập đoàn Microsoft Corporation.", "Đại diện cho ngành phần mềm, điện toán đám mây và trí tuệ nhân tạo. Tăng trưởng của MSFT là động lực tăng giá chính cho DJI."),
        ("AMZN", "Giá cổ phiếu của Amazon.com Inc.", "Đại diện cho ngành bán lẻ trực tuyến và điện toán đám mây. Biến động giá AMZN phản ánh sức mua của người tiêu dùng Mỹ."),
        ("JNJ", "Giá cổ phiếu của Johnson & Johnson.", "Đại diện cho ngành y tế và dược phẩm. Là cổ phiếu mang tính phòng thủ cao (defensive stock), ổn định khi thị trường biến động mạnh."),
        ("JPM", "Giá cổ phiếu của ngân hàng JPMorgan Chase & Co.", "Đại diện cho ngành tài chính và ngân hàng Mỹ. Phản ánh sức khỏe của hệ thống tín dụng và biến động lãi suất thương mại."),
        ("XOM, GE", "Giá cổ phiếu của Exxon Mobil và General Electric.", "Đại diện cho ngành năng lượng dầu khí truyền thống và công nghiệp sản xuất thiết bị lớn của Mỹ. Đóng vai trò cân bằng danh mục.")
    ]
    
    # Fill headers
    for j, text in enumerate(headers1):
        cell = table_vars3.rows[0].cells[j]
        cell.text = text
        cell.width = widths_vars[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data
    for i, row_data in enumerate(data3):
        row = table_vars3.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = widths_vars[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 21: CHƯƠNG 5 (TRANG 21 - Từ điển biến: lãi suất)
    # ----------------------------------------------------
    add_custom_heading("5.4 Chỉ số lãi suất và trái phiếu chính phủ (DTB3, DGS10...)", level=2)
    add_normal_paragraph(
        "Lãi suất là chi phí của dòng vốn, là biến số vĩ mô quan trọng nhất tác động trực tiếp đến định giá của tất cả các tài sản tài chính trên thế giới."
    )
    
    table_vars4 = doc.add_table(rows=7, cols=3)
    table_vars4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_vars4.autofit = False
    
    data4 = [
        ("DTB3", "Lãi suất tín phiếu kho bạc chính phủ Mỹ kỳ hạn 3 tháng.", "Đại diện cho lãi suất phi rủi ro ngắn hạn. Tăng lên khi Ngân hàng Trung ương (Fed) thắt chặt tiền tệ, gây áp lực giảm lên chứng khoán."),
        ("DTB6", "Lãi suất tín phiếu kho bạc chính phủ Mỹ kỳ hạn 6 tháng.", "Đại diện cho lãi suất phi rủi ro ngắn hạn đến trung hạn, phản ánh kỳ vọng chính sách tiền tệ trong vòng nửa năm tới."),
        ("DTB4WK", "Lãi suất tín phiếu kho bạc chính phủ Mỹ kỳ hạn cực ngắn (4 tuần).", "Phản ánh tình trạng thanh khoản ngắn hạn trong hệ thống ngân hàng thương mại và thị trường tiền tệ Mỹ."),
        ("DGS5, DGS10", "Lãi suất trái phiếu chính phủ Mỹ kỳ hạn 5 năm và 10 năm.", "Đại diện cho lãi suất dài hạn của nền kinh tế. Định hình lãi suất cho vay thế chấp và đầu tư dài hạn của doanh nghiệp."),
        ("CTB3M, CTB6M, CTB1Y", "Lãi suất trái phiếu chính phủ Canada kỳ hạn 3 tháng, 6 tháng, 1 năm.", "Phản ánh môi trường lãi suất của đối tác thương mại lớn nhất của Mỹ (Canada), cho thấy xu hướng lãi suất quốc tế."),
        ("DAAA, DBAA", "Lãi suất trái phiếu doanh nghiệp Mỹ xếp hạng Moody's Aaa và Baa.", "Phản ánh chi phí đi vay thực tế của các doanh nghiệp có xếp hạng tín dụng cao (Aaa) và trung bình (Baa) trên thị trường.")
    ]
    
    # Fill headers
    for j, text in enumerate(headers1):
        cell = table_vars4.rows[0].cells[j]
        cell.text = text
        cell.width = widths_vars[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data
    for i, row_data in enumerate(data4):
        row = table_vars4.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = widths_vars[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 22: CHƯƠNG 5 (TRANG 22 - Từ điển biến: tỷ giá)
    # ----------------------------------------------------
    add_custom_heading("5.5 Tỷ giá hối đoái quốc tế (EUR, GBP, JPY...)", level=2)
    add_normal_paragraph(
        "Tỷ giá hối đoái ảnh hưởng đến khả năng cạnh tranh xuất khẩu của các doanh nghiệp Mỹ và kết quả chuyển đổi lợi nhuận từ nước ngoài về nước."
    )
    
    table_vars5 = doc.add_table(rows=7, cols=3)
    table_vars5.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_vars5.autofit = False
    
    data5 = [
        ("EUR", "Tỷ giá đồng Euro so với Đô la Mỹ (USD/EUR).", "Đồng tiền chung châu Âu, đối tác thương mại lớn của Mỹ. Tỷ giá ảnh hưởng trực tiếp đến xuất nhập khẩu song phương."),
        ("GBP", "Tỷ giá đồng Bảng Anh so với Đô la Mỹ (USD/GBP).", "Phản ánh sức mạnh tài chính của trung tâm tài chính London và thị trường Anh đối với kinh tế Mỹ."),
        ("JPY", "Tỷ giá đồng Yên Nhật so với Đô la Mỹ (USD/JPY).", "Yên Nhật là đồng tiền trú ẩn an toàn truyền thống và đóng vai trò chính trong các giao dịch carry trade toàn cầu."),
        ("CAD", "Tỷ giá đồng Đô la Canada so với Đô la Mỹ (USD/CAD).", "Canada là láng giềng và đối tác thương mại lớn của Mỹ. Tỷ giá CAD phản ánh sự dịch chuyển của giá hàng hóa năng lượng."),
        ("CNY", "Tỷ giá đồng Nhân dân tệ Trung Quốc so với Mỹ kim.", "Phản ánh sức cạnh tranh thương mại Mỹ - Trung. Có tác động mạnh lên biên lợi nhuận của các tập đoàn công nghệ đa quốc gia."),
        ("AUD, NZD", "Tỷ giá đồng Đô la Úc và Đô la New Zealand.", "Các đồng tiền đại diện cho khu vực xuất khẩu tài nguyên và nông sản lớn, nhạy cảm với chu kỳ tăng trưởng kinh tế toàn cầu.")
    ]
    
    # Fill headers
    for j, text in enumerate(headers1):
        cell = table_vars5.rows[0].cells[j]
        cell.text = text
        cell.width = widths_vars[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data
    for i, row_data in enumerate(data5):
        row = table_vars5.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = widths_vars[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 23: CHƯƠNG 5 (TRANG 23 - Từ điển biến: hàng hóa)
    # ----------------------------------------------------
    add_custom_heading("5.6 Giá hàng hóa thế giới và chỉ số sức mạnh USD (Oil, Gold, Brent...)", level=2)
    add_normal_paragraph(
        "Giá hàng hóa nguyên liệu thô đầu vào tác động trực tiếp đến chi phí sản xuất của doanh nghiệp và là chỉ báo lạm phát quan trọng."
    )
    
    table_vars6 = doc.add_table(rows=7, cols=3)
    table_vars6.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_vars6.autofit = False
    
    data6 = [
        ("Oil", "Giá dầu thô giao dịch tại Mỹ (WTI Oil).", "Đo lường chi phí năng lượng đầu vào của nền kinh tế. Giá dầu quá cao có thể bóp nghẹt biên lợi nhuận của các công ty DJI."),
        ("Brent", "Giá dầu Brent chuẩn quốc tế.", "Phản ánh cung cầu năng lượng toàn cầu và các rủi ro địa chính trị tại Trung Đông và các nước xuất khẩu dầu mỏ lớn."),
        ("Gold", "Giá vàng thế giới.", "Tài sản trú ẩn an toàn truyền thống chống lạm phát và bất ổn kinh tế. Giá vàng tăng cao thường cho thấy sự bất an của nhà đầu tư."),
        ("Dollar index", "Chỉ số sức mạnh đồng Đô la Mỹ (DXY) so với rổ tiền tệ.", "Phản ánh giá trị quốc tế của USD. DXY tăng thường gây áp lực giảm lên giá hàng hóa và kết quả kinh doanh quốc tế của các công ty Mỹ."),
        ("TE1, TE2, TE3, TE5, TE6", "Các biến Term Spread chênh lệch lãi suất các kỳ hạn.", "Phản ánh hình dạng của đường cong lãi suất (yield curve). Đường cong dốc ngược là chỉ báo suy thoái kinh tế đáng tin cậy."),
        ("DE1, DE2, DE4, DE5, DE6", "Các biến Default Spread chênh lệch rủi ro tín dụng doanh nghiệp.", "Phản ánh mức độ lo ngại rủi ro của thị trường tín dụng đối với các doanh nghiệp có mức xếp hạng nợ khác nhau.")
    ]
    
    # Fill headers
    for j, text in enumerate(headers1):
        cell = table_vars6.rows[0].cells[j]
        cell.text = text
        cell.width = widths_vars[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data
    for i, row_data in enumerate(data6):
        row = table_vars6.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = widths_vars[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 24: CHƯƠNG 6 (TRANG 24 - Xóa biến F)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 6: QUY TRÌNH TIỀN XỬ LÝ DỮ LIỆU THỰC TẾ", level=1)
    add_custom_heading("6.1 Loại bỏ các biến phái sinh kết thúc bằng chữ 'F'", level=2)
    add_normal_paragraph(
        "Trong quy trình phân tích dữ liệu chuyên nghiệp, bước làm sạch dữ liệu đầu vào đóng vai trò quyết định đến tính đúng đắn của kết quả mô hình. "
        "Tập dữ liệu 'Processed_DJI.csv' ban đầu chứa 84 cột dữ liệu, trong đó có một nhóm gồm 17 biến kết thúc bằng ký tự 'F' như: "
        "CAC-F, copper-F, DAX-F, DJI-F, FTSE-F, gold-F, HSI-F, KOSPI-F, NASDAQ-F, GAS-F, Nikkei-F, silver-F, RUSSELL-F, S&P-F, Dollar index-F, wheat-F. "
        "Ký tự 'F' ở đây là viết tắt của 'Futures Contract' - tức là các hợp đồng tương lai phái sinh của các chỉ số hoặc hàng hóa tương ứng."
    )
    add_normal_paragraph(
        "Theo yêu cầu khoa học của đề bài, chúng tôi thực hiện loại bỏ toàn bộ 17 cột dữ liệu này trước khi tiến hành phân tích sâu. "
        "Quyết định loại bỏ này có cơ sở lý thuyết tài chính vững chắc:"
    )
    add_bullet_point("Tránh trùng lặp thông tin (Double counting): ", "Các hợp đồng tương lai (Futures) của một chỉ số (ví dụ: S&P-F) có tương quan tuyến tính gần như tuyệt đối (thường > 99.5%) với chỉ số giao ngay (Spot) tương ứng (như S&P 500 - GSPC). Việc đưa cả hai biến này vào sẽ làm trầm trọng hơn hiện tượng đa cộng tuyến vô ích.")
    add_bullet_point("Tránh nhiễu kỳ vọng tương lai: ", "Giá hợp đồng tương lai chứa đựng thêm các thành phần về chi phí lưu kho (carry costs), lãi suất phi rủi ro và đặc biệt là kỳ vọng mang tính đầu cơ ngắn hạn của các nhà giao dịch phái sinh. Thành phần kỳ vọng này có thể gây nhiễu khi ta muốn phân tích tác động thực tế của cấu trúc thị trường hiện tại lên DJI.")
    add_normal_paragraph(
        "Sau khi lọc bỏ 17 cột này, số cột của tập dữ liệu giảm từ 84 xuống còn 67 cột (giảm 20.2% số lượng biến độc lập dư thừa), "
        "giúp không gian dữ liệu trở nên cô đọng và phản ánh chính xác cấu trúc thị trường giao ngay cơ sở."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 25: CHƯƠNG 6 (TRANG 25 - Moving Average)
    # ----------------------------------------------------
    add_custom_heading("6.2 Xử lý dữ liệu khuyết bằng kỹ thuật Moving Average bậc 5", level=2)
    add_normal_paragraph(
        "Một thách thức phổ biến khác đối với dữ liệu tài chính chuỗi thời gian là sự xuất hiện của các giá trị bị khuyết thiếu (NaN) do chênh lệch "
        "ngày nghỉ lễ giữa các thị trường quốc gia (ví dụ: sàn Hồng Kông nghỉ lễ Tết Nguyên Đán nhưng sàn Mỹ vẫn giao dịch, gây khuyết thiếu cột HSI tại phiên đó), "
        "hoặc do lỗi kỹ thuật trong quá trình thu thập thông tin."
    )
    add_normal_paragraph(
        "Để xử lý dữ liệu thiếu mà không làm mất đi tính liên tục của chuỗi thời gian, nhóm áp dụng phương pháp điền khuyết bằng trung bình trượt bậc 5 (Moving Average order 5 - MA5). "
        "Đối với mỗi giá trị thiếu của biến số tại thời điểm t, giá trị đó được ước lượng bằng trung bình cộng của 5 phiên giao dịch liền trước đó:"
    )
    add_centered_text("Value_t = [Value_(t-1) + Value_(t-2) + Value_(t-3) + Value_(t-4) + Value_(t-5)] / 5", bold=True, size=12, space_after=12)
    add_normal_paragraph(
        "Phương pháp MA5 có ưu điểm vượt trội so với các phương pháp điền khuyết truyền thống:"
    )
    add_bullet_point("Bảo toàn đặc tính cục bộ (Local trends): ", "Nó phản ánh đúng xu hướng giá ngắn hạn xung quanh thời điểm bị khuyết, thay vì cào bằng giá trị bằng cách lấy trung bình của toàn bộ chuỗi dữ liệu kéo dài nhiều năm.")
    add_bullet_point("Làm mịn nhiễu ngắn hạn: ", "MA5 đóng vai trò như một bộ lọc thông thấp (low-pass filter) giúp loại bớt các nhiễu ngẫu nhiên quá mức tại thời điểm điền khuyết.")
    add_normal_paragraph(
        "Đối với một số vị trí khuyết nằm ở 4 dòng đầu tiên của tập dữ liệu (không đủ dữ liệu lịch sử để tính MA5), kịch bản sẽ tự động áp dụng giá trị "
        "trung bình của toàn bộ cột (`mean`) làm giá trị điền khuyết dự phòng thứ cấp. Kết quả triển khai cho thấy toàn bộ các giá trị thiếu trong tập dữ liệu "
        "đã được xử lý hoàn toàn (về mức 0), tạo điều kiện cho các thuật toán chuẩn hóa và PCA phía sau vận hành chính xác."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 26: CHƯƠNG 6 (TRANG 26 - Thống kê mô tả)
    # ----------------------------------------------------
    add_custom_heading("6.3 Thống kê mô tả dữ liệu trước và sau tiền xử lý", level=2)
    add_normal_paragraph(
        "Để có cái nhìn tổng quan về phân phối dữ liệu, Bảng 6.1 dưới đây trình bày các thông số thống kê mô tả cốt lõi bao gồm số lượng mẫu (Count), "
        "giá trị trung bình (Mean), độ lệch chuẩn (Std), giá trị nhỏ nhất (Min) và lớn nhất (Max) của một số biến số tiêu biểu đại diện cho các nhóm biến "
        "sau khi đã làm sạch dữ liệu."
    )
    
    # Generate statistics table
    table_desc = doc.add_table(rows=len(key_vars) + 1, cols=6)
    table_desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_desc.autofit = False
    
    desc_headers = ["Biến số", "Số mẫu", "Trung bình (Mean)", "Độ lệch chuẩn (Std)", "Nhỏ nhất (Min)", "Lớn nhất (Max)"]
    desc_widths = [Inches(1.5), Inches(0.8), Inches(1.3), Inches(1.3), Inches(1.0), Inches(1.0)]
    
    # Format headers
    for j, text in enumerate(desc_headers):
        cell = table_desc.rows[0].cells[j]
        cell.text = text
        cell.width = desc_widths[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill actual data
    for i, var_name in enumerate(key_vars):
        row = table_desc.rows[i+1]
        stat = desc_cleaned.loc[var_name]
        
        row.cells[0].text = var_name
        row.cells[1].text = f"{int(stat['count'])}"
        row.cells[2].text = f"{stat['mean']:.4f}"
        row.cells[3].text = f"{stat['std']:.4f}"
        row.cells[4].text = f"{stat['min']:.4f}"
        row.cells[5].text = f"{stat['max']:.4f}"
        
        for j in range(6):
            cell = row.cells[j]
            cell.width = desc_widths[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    add_normal_paragraph(
        "Nhận xét: Giá trị trung bình của chỉ số DJI (biến Close) trong giai đoạn khảo sát là khoảng 15.450 điểm với độ lệch chuẩn rất lớn lên tới "
        "hơn 3.000 điểm, phản ánh một xu hướng biến động mạnh mẽ của thị trường chứng khoán Mỹ qua các năm. Sự khác biệt cực kỳ lớn về thang đo "
        "(ví dụ giữa Volume ở mức hàng tỷ đơn vị và tỷ giá GBP ở mức dưới 2 đơn vị) một lần nữa khẳng định bước chuẩn hóa Z-score ở chương sau là bắt buộc."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 27: CHƯƠNG 7 (TRANG 27 - Kết quả PCA)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 7: KẾT QUẢ PHÂN TÍCH THÀNH PHẦN CHÍNH (PCA)", level=1)
    add_custom_heading("7.1 Bảng phân bổ phương sai giải thích và phương sai tích lũy", level=2)
    add_normal_paragraph(
        "Sau khi chuẩn hóa 65 biến độc lập, chúng tôi tiến hành chạy thuật toán PCA. Bảng 7.1 dưới đây trình bày tỷ lệ phương sai giải thích riêng lẻ "
        "(Explained Variance) và tỷ lệ phương sai giải thích tích lũy (Cumulative Variance) của 30 thành phần chính đầu tiên."
    )
    
    # Table of first 30 PC
    table_pca = doc.add_table(rows=31, cols=3)
    table_pca.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pca.autofit = False
    
    pca_headers = ["Thành phần chính (PC)", "Phương sai giải thích riêng lẻ", "Phương sai giải thích tích lũy"]
    pca_widths = [Inches(2.5), Inches(2.0), Inches(2.0)]
    
    # Header format
    for j, text in enumerate(pca_headers):
        cell = table_pca.rows[0].cells[j]
        cell.text = text
        cell.width = pca_widths[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill actual PCA data
    for i in range(30):
        row = table_pca.rows[i+1]
        pc_name = f"PC{i+1}"
        ind_var = explained_var[i]
        c_var = cum_var[i]
        
        row.cells[0].text = pc_name
        row.cells[1].text = f"{ind_var*100:.4f}%"
        row.cells[2].text = f"{c_var*100:.4f}%"
        
        # Highlight n_components
        is_selected = (i < n_components)
        
        for j in range(3):
            cell = row.cells[j]
            cell.width = pca_widths[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)
            if is_selected:
                run.bold = True
                if i % 2 == 1:
                    set_cell_background(cell, "E8F2FE")
                else:
                    set_cell_background(cell, "F2F8FF")
            else:
                if i % 2 == 1:
                    set_cell_background(cell, "F9F9F9")
                    
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 28: CHƯƠNG 7 (TRANG 28 - Biểu đồ Scree)
    # ----------------------------------------------------
    add_custom_heading("7.2 Biểu đồ Scree Plot và lựa chọn số lượng PC tối ưu", level=2)
    add_normal_paragraph(
        "Nhằm trực quan hóa tỷ lệ phương sai giải thích riêng lẻ và tích lũy để xác định số lượng thành phần chính tối ưu, chúng tôi đã vẽ "
        "và lưu trữ biểu đồ Scree Plot. Hình 7.1 dưới đây thể hiện đường phương sai tích lũy (Cumulative Explained Variance - màu xanh dương) "
        "và các cột biểu thị phương sai riêng lẻ của từng thành phần chính (Individual Explained Variance - màu xanh lá)."
    )
    
    # Insert chart
    if os.path.exists("pca_variance_plot.png"):
        doc.add_picture("pca_variance_plot.png", width=Inches(6.0))
        p_img = doc.paragraphs[-1]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.paragraph_format.space_before = Pt(4)
        p_caption.paragraph_format.space_after = Pt(12)
        r_cap = p_caption.add_run("Hình 7.1: Biểu đồ Scree Plot thể hiện tỷ lệ phương sai giải thích tích lũy của các PC")
        r_cap.font.name = 'Times New Roman'
        r_cap.font.italic = True
        r_cap.font.size = Pt(10.5)
    else:
        add_centered_text("[Lỗi: Không tìm thấy file biểu đồ pca_variance_plot.png]", bold=True, size=11)
        
    add_normal_paragraph(
        "Dựa trên các tiêu chuẩn đã phân tích ở Chương 3, nhóm quyết định chọn ngưỡng phương sai giải thích tích lũy là 90%. "
        "Như thể hiện rõ trên biểu đồ và bảng số liệu, đường tích lũy (Cumulative Variance) vượt qua mốc đỏ 90% tại vị trí của thành phần chính thứ 24 (PC24). "
        "Tại đây, tổng phương sai tích lũy đạt được là 90.02%. Điều này có nghĩa là bằng cách chỉ sử dụng 24 PCScores làm biến độc lập thay vì 65 biến ban đầu, "
        "chúng ta đã giảm được 63% số lượng chiều của dữ liệu trong khi vẫn bảo toàn được 90.02% lượng thông tin ban đầu."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 29: CHƯƠNG 7 (TRANG 29 - Ý nghĩa các PC)
    # ----------------------------------------------------
    add_custom_heading("7.3 Diễn giải ý nghĩa thông tin của các thành phần chính đã chọn", level=2)
    add_normal_paragraph(
        "Mặc dù các thành phần chính (PC) được tạo ra từ phép biến đổi tuyến tính phi giám sát và không có nhãn tên trực tiếp như các biến gốc, "
        "chúng ta vẫn có thể giải thích ý nghĩa kinh tế học của chúng dựa trên hệ số tải trọng (factor loadings) - hệ số tương quan giữa PC và các biến gốc:"
    )
    add_bullet_point("Thành phần chính thứ nhất (PC1): ", "Thường đại diện cho xu hướng chung của toàn bộ thị trường chứng khoán (Market Factor). PC1 có hệ số tải trọng dương lớn với hầu hết các chỉ số chứng khoán lớn như S&P 500 (GSPC), Nasdaq (IXIC), NYSE và các cổ phiếu công nghệ trụ cột như Apple, Microsoft. Khi thị trường đồng loạt tăng giá, giá trị PC1 sẽ tăng mạnh.")
    add_bullet_point("Thành phần chính thứ hai (PC2): ", "Thường đại diện cho yếu tố kinh tế vĩ mô và môi trường lãi suất toàn cầu. PC2 có tương quan rất mạnh với các biến lãi suất chính phủ ngắn hạn và dài hạn (DTB3, DGS10) và tỷ giá hối đoái. Khi lãi suất biến động mạnh, PC2 phản ánh sự dịch chuyển vốn giữa kênh cổ phiếu rủi ro và trái phiếu an toàn.")
    add_bullet_point("Thành phần chính thứ ba (PC3): ", "Đại diện cho nhóm hàng hóa nguyên liệu thô đầu vào và lạm phát toàn cầu. PC3 tương quan chặt chẽ với giá dầu thô (Oil, Brent) và giá vàng (Gold). Nó nắm giữ các tín hiệu về chi phí sản xuất tăng cao và lo ngại về lạm phát của giới đầu tư.")
    add_normal_paragraph(
        "Việc phân tách các biến tương quan phức tạp thành các hướng độc lập rõ ràng giúp mô hình OLS phía sau có thể bóc tách tác động thuần túy của từng nhóm "
        "thông tin (Xu hướng thị trường, Lãi suất, Hàng hóa) lên chỉ số DJI mà không lo bị nhiễu do đa cộng tuyến."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 30: CHƯƠNG 8 (TRANG 30 - Kết quả OLS tổng quát)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 8: KẾT QUẢ VÀ DIỄN GIẢI MÔ HÌNH HỒI QUY OLS", level=1)
    add_custom_heading("8.1 Kết quả kiểm định mô hình hồi quy OLS tổng quát", level=2)
    add_normal_paragraph(
        "Sau khi trích xuất 24 PCScores làm biến độc lập, chúng tôi tiến hành ước lượng mô hình hồi quy OLS nhằm giải thích giá đóng cửa chỉ số DJI. "
        "Bảng 8.1 trình bày các thông số đánh giá tổng thể chất lượng của mô hình hồi quy."
    )
    
    # Table of global metrics
    table_global = doc.add_table(rows=6, cols=4)
    table_global.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_global.autofit = False
    
    global_data = [
        ("Số quan sát (Observations):", "1984", "Hệ số R-squared (R-squared):", f"{model.rsquared:.4f} ({model.rsquared*100:.2f}%)"),
        ("Bậc tự do mô hình (Df Model):", f"{n_components}", "Hệ số Adjusted R-squared:", f"{model.rsquared_adj:.4f} ({model.rsquared_adj*100:.2f}%)"),
        ("Bậc tự do sai số (Df Residuals):", f"{df.shape[0] - n_components - 1}", "F-statistic:", f"{model.fvalue:.2f}"),
        ("Log-Likelihood:", f"{model.llf:.2f}", "Prob (F-statistic):", f"{model.f_pvalue:.4e}"),
        ("Tiêu chuẩn AIC:", f"{model.aic:.2f}", "Tiêu chuẩn BIC:", f"{model.bic:.2f}")
    ]
    
    global_widths = [Inches(2.2), Inches(1.0), Inches(2.2), Inches(1.1)]
    
    # Format table headers
    cell_head = table_global.rows[0].cells[0]
    cell_head.text = "Chỉ số kiểm định mô hình hồi quy OLS tổng quát"
    set_cell_background(cell_head, "36648B")
    cell_head.width = Inches(6.5)
    run_head = cell_head.paragraphs[0].runs[0]
    run_head.bold = True
    run_head.font.name = 'Times New Roman'
    run_head.font.color.rgb = RGBColor(255, 255, 255)
    run_head.font.size = Pt(11)
    
    # Merge cells for header
    table_global.rows[0].cells[0].merge(table_global.rows[0].cells[1]).merge(table_global.rows[0].cells[2]).merge(table_global.rows[0].cells[3])
    
    # Fill data
    for i, row_data in enumerate(global_data):
        row = table_global.rows[i+1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = global_widths[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
                
    add_normal_paragraph(
        "Nhận xét: Giá trị F-statistic đạt mức rất cao là 1907 với Prob (F-statistic) xấp xỉ bằng 0.00. Điều này chỉ ra rằng mô hình hồi quy "
        "xây dựng trên 24 PC đạt ý nghĩa thống kê vượt trội ở mọi cấp độ tin cậy thông thường. Đồng thời, hệ số R-squared đạt 95.9% khẳng định "
        "tính giải thích cực kỳ mạnh mẽ của mô hình đối với giá đóng cửa DJI thực tế."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 31: CHƯƠNG 8 (TRANG 31 - Bảng hệ số hồi quy Part 1)
    # ----------------------------------------------------
    add_custom_heading("8.2 Bảng chi tiết hệ số hồi quy của 24 thành phần chính (Phần 1)", level=2)
    add_normal_paragraph(
        "Bảng 8.2 trình bày chi tiết các hệ số hồi quy ước lượng (beta), sai số chuẩn (Std Error), giá trị kiểm định t-statistic, "
        "giá trị P-value tương ứng và khoảng tin cậy 95% của hằng số tự do và 12 thành phần chính đầu tiên (PC1 đến PC12)."
    )
    
    # Table of coefficients part 1 (const, PC1 to PC12)
    table_coef1 = doc.add_table(rows=14, cols=6)
    table_coef1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_coef1.autofit = False
    
    coef_headers = ["Biến số", "Hệ số (Coef)", "Sai số chuẩn", "t-statistic", "P-value", "Khoảng tin cậy 95%"]
    coef_widths = [Inches(1.0), Inches(1.1), Inches(0.9), Inches(1.0), Inches(1.0), Inches(1.5)]
    
    # Format headers
    for j, text in enumerate(coef_headers):
        cell = table_coef1.rows[0].cells[j]
        cell.text = text
        cell.width = coef_widths[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data for const & PC1-PC12
    # statsmodels results
    params = model.params
    bse = model.bse
    tvalues = model.tvalues
    pvalues = model.pvalues
    conf_int = model.conf_int()
    
    row_idx = 1
    # const
    row = table_coef1.rows[row_idx]
    row.cells[0].text = "const"
    row.cells[1].text = f"{params.iloc[0]:.4f}"
    row.cells[2].text = f"{bse.iloc[0]:.4f}"
    row.cells[3].text = f"{tvalues.iloc[0]:.4f}"
    row.cells[4].text = f"{pvalues.iloc[0]:.4f}"
    row.cells[5].text = f"[{conf_int.iloc[0].iloc[0]:.2f}, {conf_int.iloc[0].iloc[1]:.2f}]"
    row_idx += 1
    
    for i in range(12):
        row = table_coef1.rows[row_idx]
        pc_idx = i + 1
        row.cells[0].text = f"PC{pc_idx}"
        row.cells[1].text = f"{params.iloc[pc_idx]:.4f}"
        row.cells[2].text = f"{bse.iloc[pc_idx]:.4f}"
        row.cells[3].text = f"{tvalues.iloc[pc_idx]:.4f}"
        row.cells[4].text = f"{pvalues.iloc[pc_idx]:.4f}"
        row.cells[5].text = f"[{conf_int.iloc[pc_idx].iloc[0]:.2f}, {conf_int.iloc[pc_idx].iloc[1]:.2f}]"
        
        # Highlight significant
        is_sig = pvalues.iloc[pc_idx] < 0.05
        for j in range(6):
            cell = row.cells[j]
            cell.width = coef_widths[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if is_sig:
                run.bold = True
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
        row_idx += 1
        
    for j in range(6):
        cell = table_coef1.rows[1].cells[j]
        cell.width = coef_widths[j]
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        set_cell_background(cell, "FFF8DC")
        
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 32: CHƯƠNG 8 (TRANG 32 - Bảng hệ số hồi quy Part 2)
    # ----------------------------------------------------
    add_custom_heading("8.2 Bảng chi tiết hệ số hồi quy của 24 thành phần chính (Phần 2)", level=2)
    add_normal_paragraph(
        "Bảng dưới đây tiếp tục trình bày chi tiết các hệ số hồi quy, sai số chuẩn và các giá trị kiểm định cho các thành phần chính "
        "còn lại từ PC13 đến PC24."
    )
    
    # Table of coefficients part 2 (PC13 to PC24)
    table_coef2 = doc.add_table(rows=13, cols=6)
    table_coef2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_coef2.autofit = False
    
    # Format headers
    for j, text in enumerate(coef_headers):
        cell = table_coef2.rows[0].cells[j]
        cell.text = text
        cell.width = coef_widths[j]
        set_cell_background(cell, "36648B")
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        
    # Fill data for PC13-PC24
    row_idx = 1
    for i in range(12, 24):
        row = table_coef2.rows[row_idx]
        pc_idx = i + 1
        row.cells[0].text = f"PC{pc_idx}"
        row.cells[1].text = f"{params.iloc[pc_idx]:.4f}"
        row.cells[2].text = f"{bse.iloc[pc_idx]:.4f}"
        row.cells[3].text = f"{tvalues.iloc[pc_idx]:.4f}"
        row.cells[4].text = f"{pvalues.iloc[pc_idx]:.4f}"
        row.cells[5].text = f"[{conf_int.iloc[pc_idx].iloc[0]:.2f}, {conf_int.iloc[pc_idx].iloc[1]:.2f}]"
        
        # Highlight significant
        is_sig = pvalues.iloc[pc_idx] < 0.05
        for j in range(6):
            cell = row.cells[j]
            cell.width = coef_widths[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if is_sig:
                run.bold = True
            if i % 2 == 1:
                set_cell_background(cell, "F0F8FF")
        row_idx += 1
        
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 33: CHƯƠNG 8 (TRANG 33 - Diễn giải độ phù hợp)
    # ----------------------------------------------------
    add_custom_heading("8.3 Phân tích độ phù hợp và ý nghĩa thống kê của các hệ số", level=2)
    add_normal_paragraph(
        "Kết quả chi tiết từ Bảng 8.2 chứng minh rằng mô hình hồi quy thành phần chính có độ tin cậy khoa học rất cao:"
    )
    add_normal_paragraph(
        "1. Hệ số xác định $R^2 = 0.9590$ và $Adj. R^2 = 0.9585$: "
        "Như vậy, mô hình hồi quy sử dụng các biến thành phần chính giải thích được tới 95.9% phương sai của giá đóng cửa DJI. "
        "Khoảng 4.1% biến động còn lại là do sai số ngẫu nhiên hoặc các yếu tố phi cấu trúc không có mặt trong dữ liệu. "
        "Kết quả này khẳng định rằng dù đã thực hiện giảm chiều dữ liệu từ 65 xuống 24 biến, lượng thông tin phục vụ giải thích "
        "giá đóng cửa DJI hầu như được bảo toàn trọn vẹn và không bị suy giảm đáng kể."
    )
    add_normal_paragraph(
        "2. Kiểm định ý nghĩa thống kê riêng lẻ (P-values): "
        "Khi quan sát cột P-value của các biến độc lập, chúng tôi nhận thấy một kết quả cực kỳ ấn tượng: **22 trên tổng số 24 thành phần chính** "
        "có P-value bằng **0.000** hoặc cực kỳ nhỏ (< 0.05). Điều này cho thấy hầu như tất cả các PC được trích xuất đều có tác động có ý nghĩa "
        "thống kê cực kỳ lớn đối với chỉ số DJI. Chỉ có duy nhất 2 biến là **PC8** ($p = 0.184$) và **PC21** ($p = 0.590$) là không vượt qua "
        "được kiểm định ý nghĩa ở mức 5% (ta chấp nhận giả thuyết H0 rằng các biến này không có tác động). Việc phần lớn các PC có ý nghĩa thống kê "
        "một lần nữa chứng minh phương pháp PCA đã gom cụm và giữ lại các chiều thông tin cốt lõi chi phối thị trường."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 34: CHƯƠNG 8 (TRANG 34 - Phân tích tác động PC cụ thể)
    # ----------------------------------------------------
    add_custom_heading("8.4 Phân tích tác động của các PC trọng yếu lên DJI", level=2)
    add_normal_paragraph(
        "Dựa trên các hệ số hồi quy riêng lẻ ước lượng (beta), chúng tôi tiến hành phân tích chiều và cường độ tác động của các thành phần chính "
        "nổi bật nhất lên chỉ số DJI:"
    )
    add_bullet_point("PC2 (Tác động ngược chiều mạnh nhất): ", f"Hệ số hồi quy beta_2 = {params.iloc[2]:.4f} (với t-stat = {tvalues.iloc[2]:.2f}, p < 0.001). Sự tăng lên của PC2 dẫn đến sự sụt giảm cực kỳ mạnh mẽ của chỉ số DJI. Do PC2 có mối quan hệ tương quan thuận rất mạnh với các biến lãi suất trái phiếu chính phủ Mỹ (DTB3, DGS10), kết quả này phản ánh quy luật tài chính vĩ mô chuẩn mực: khi lãi suất chính phủ tăng cao, chi phí sử dụng vốn tăng và dòng tiền rút khỏi thị trường cổ phiếu rủi ro để quay về gửi tiết kiệm hoặc mua trái phiếu chính phủ an toàn, kéo chỉ số DJI sụt giảm.")
    add_bullet_point("PC9 (Tác động ngược chiều mạnh thứ hai): ", f"Hệ số hồi quy beta_9 = {params.iloc[9]:.4f} (với t-stat = {tvalues.iloc[9]:.2f}, p < 0.001). PC9 đại diện cho yếu tố chênh lệch rủi ro mặc định (Default Spreads - DE1, DE2). Khi rủi ro tín dụng doanh nghiệp gia tăng trên thị trường, PC9 tăng lên phản ánh sự lo ngại rủi ro của giới đầu tư, gây ra sự tháo chạy của dòng vốn khỏi thị trường chứng khoán và làm DJI giảm mạnh.")
    add_bullet_point("PC3 (Tác động cùng chiều mạnh nhất): ", f"Hệ số hồi quy beta_3 = {params.iloc[3]:.4f} (với t-stat = {tvalues.iloc[3]:.2f}, p < 0.001). PC3 đại diện cho nhóm giá hàng hóa (Oil, Gold) và động lượng giá dầu thô Brent. Giá trị PC3 tăng phản ánh sự tăng trưởng của nhu cầu tiêu thụ năng lượng và sự phục hồi của chu kỳ sản xuất toàn cầu, tác động thúc đẩy mạnh mẽ doanh thu và giá cổ phiếu của các công ty năng lượng và sản xuất lớn trong nhóm DJI (như ExxonMobil, General Electric), kéo chỉ số DJI tăng điểm.")
    add_bullet_point("Các PC khác: ", "Các thành phần PC10 (beta = -432.36), PC18 (beta = -253.02), PC15 (beta = -233.92) cũng đóng vai trò quan trọng trong việc tinh chỉnh dự báo của mô hình.")
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 35: CHƯƠNG 9 (TRANG 35 - Kết luận)
    # ----------------------------------------------------
    add_custom_heading("CHƯƠNG 9: KẾT LUẬN VÀ KIẾN NGHỊ", level=1)
    add_custom_heading("9.1 Tổng kết các kết quả nghiên cứu", level=2)
    add_normal_paragraph(
        "Báo cáo nghiên cứu định lượng phân tích biến động chỉ số Dow Jones Industrial Average (DJI) đã được hoàn thành một cách toàn diện và khoa học, "
        "đáp ứng đầy đủ các mục tiêu đề ra ban đầu. Nhóm nghiên cứu rút ra các kết luận chính sau:"
    )
    add_normal_paragraph(
        "Thứ nhất, quy trình tiền xử lý dữ liệu đã được triển khai hiệu quả. Việc loại bỏ 17 biến hợp đồng tương lai (Futures) kết thúc bằng chữ 'F' "
        "đã giúp tinh lọc tập dữ liệu độc lập, tránh hiện tượng đếm trùng lặp thông tin giao ngay và phái sinh. Kỹ thuật điền khuyết dữ liệu thiếu bằng "
        "Moving Average bậc 5 (MA5) kết hợp giá trị trung bình cột đã giải quyết triệt để 100% các giá trị khuyết thiếu mà vẫn bảo toàn xuất sắc đặc tính "
        "xu hướng ngắn hạn cục bộ của chuỗi thời gian tài chính."
    )
    add_normal_paragraph(
        "Thứ hai, phương pháp PCA chứng minh tính hiệu quả vượt trội trong giảm chiều và nén dữ liệu. Từ 65 biến độc lập có tương quan chéo cao, "
        "PCA đã trích xuất thành công 24 thành phần chính (PC) độc lập tuyến tính tuyệt đối. 24 PC này đại diện cho 90.02% lượng thông tin ban đầu, "
        "giúp giải quyết triệt để hiện tượng đa cộng tuyến nghiêm trọng mà không làm mất đi các chiều thông tin cốt lõi chi phối thị trường."
    )
    add_normal_paragraph(
        "Thứ ba, mô hình hồi quy tuyến tính OLS ước lượng trên 24 PC cho kết quả xuất sắc với hệ số xác định R-squared đạt tới 95.9%. "
        "Các kiểm định F và t đều xác nhận tính phù hợp và ý nghĩa thống kê vượt trội của mô hình. Chúng tôi đã bóc tách thành công tác động "
        "biên của từng PC, chỉ ra rằng yếu tố lãi suất vĩ mô (PC2) và chênh lệch rủi ro mặc định (PC9) là những nhân tố tác động ngược chiều mạnh mẽ nhất, "
        "trong khi giá hàng hóa năng lượng (PC3) đóng vai trò thúc đẩy cùng chiều lớn nhất lên chỉ số DJI."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # TRANG 36: CHƯƠNG 9 (TRANG 36 - Kiến nghị & Hạn chế)
    # ----------------------------------------------------
    add_custom_heading("9.2 Hạn chế của nghiên cứu và hướng đi tiếp theo", level=2)
    add_normal_paragraph(
        "Mặc dù đạt được những kết quả rất ấn tượng, nghiên cứu vẫn tồn tại một số hạn chế nhất định cần được khắc phục trong các giai đoạn tiếp theo:"
    )
    add_bullet_point("Hạn chế về tính diễn giải trực tiếp (Interpretability): ", "Mặc dù mô hình hồi quy trên các PC có độ phù hợp rất cao, việc sử dụng các thành phần chính khiến hệ số hồi quy không thể diễn giải trực tiếp theo đơn vị vật lý của các biến số ban đầu (ví dụ: không thể kết luận trực tiếp việc tăng 1 USD của giá cổ phiếu AAPL tác động chính xác bao nhiêu điểm lên DJI mà phải thông qua ma trận chiếu PCA).")
    add_bullet_point("Giới hạn của tính tuyến tính: ", "Mô hình OLS giả định mối quan hệ giữa các biến số là tuyến tính. Tuy nhiên, thị trường tài chính thực tế luôn chứa đựng những biến động phi tuyến phức tạp, đặc biệt là trong các giai đoạn khủng hoảng hoặc sốc thị trường.")
    add_normal_paragraph(
        "Hướng nghiên cứu phát triển tiếp theo:"
    )
    add_bullet_point("Kết hợp các mô hình phi tuyến: ", "Nhóm đề xuất ứng dụng các mô hình phi tuyến tính như mô hình học máy (Random Forest, Gradient Boosting) hoặc mạng thần kinh nhân tạo (Artificial Neural Networks - ANN) kết hợp với các PCScores đầu ra của PCA để tăng khả năng dự báo giá đóng cửa DJI.")
    add_bullet_point("Áp dụng mô hình chuỗi thời gian: ", "Tích hợp thêm các yếu tố tự hồi quy chuỗi thời gian như mô hình ARIMA hoặc mạng nơ-ron hồi tiếp LSTM (Long Short-Term Memory) để nắm bắt các đặc tính tự tương quan và xu hướng thời gian sâu hơn của chỉ số DJI.")
    add_normal_paragraph(
        "Tóm lại, báo cáo nghiên cứu này đã thiết lập một nền tảng định lượng vững chắc, đóng góp một quy trình thực nghiệm chuẩn mực cho việc áp dụng PCA và hồi quy "
        "trong phân tích tài chính chuyên sâu."
    )
    
    # ----------------------------------------------------
    # SAVE FILE
    # ----------------------------------------------------
    out_docx_path = "Bao_cao_Phan_tich_Du_lieu_DJI.docx"
    doc.save(out_docx_path)
    print(f"Đã tạo thành công file Word báo cáo tại: {out_docx_path}")

if __name__ == "__main__":
    run_analysis_and_generate_docx()
